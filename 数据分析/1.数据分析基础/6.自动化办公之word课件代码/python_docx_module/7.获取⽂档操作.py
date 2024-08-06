from docx import Document

docu = Document('test.docx')
paragraphs = docu.paragraphs
for paragraph in paragraphs:
    print(paragraph.text)

# 段落数量
print(len(docu.paragraphs))
