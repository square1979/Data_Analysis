from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn

# 新建文档对象
doc = Document()

# 全局字体设置
doc.styles['Normal'].font.name = u'宋体'
doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

# 1、标题
title = doc.add_paragraph()
# 内容
title1 = title.add_run('员工请假条')
# 设置字体⼤⼩
title1.font.size = Pt(20)
# 加粗
title1.bold = True
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 2、正⽂
article1 = doc.add_paragraph()
a1 = article1.add_run('____部:')
article2 = doc.add_paragraph()
a2 = article2.add_run('因个人原因（如家庭紧急事务/身体不适等，请根据实际情况填写具体原因）'
                      '，特向您请事假__天。请假时间自____年__月__日至____年__月__日。这段时间内原计划安排的课程/工作任务已提前做好相应的调整与安排'
                      '确保不会给团队/教学进度带来不便。恳请领导批准为盼')

# ⾸⾏缩进 负值表示悬挂缩进
article2_format = article2.paragraph_format
article2_format.first_line_indent = Inches(0.3)

# 设置请假人姓名区域
article3 = doc.add_paragraph()
a3 = article3.add_run('请假人:')  # 内容
article3.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
article3_format = article3.paragraph_format
article3_format.right_indent = Inches(0.9)

# 设置请假日期区域
nowData = doc.add_paragraph()
n3 = nowData.add_run('年 ⽉ ⽇')  # 内容
nowData.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
nowData_format = nowData.paragraph_format
nowData_format.right_indent = Inches(0.3)

doc.save('test.docx')
