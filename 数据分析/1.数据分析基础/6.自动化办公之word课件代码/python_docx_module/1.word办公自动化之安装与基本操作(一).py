from docx import Document
from docx.shared import Pt

doc = Document()

# 添加标题
doc.add_heading('总标题', 0)
doc.add_heading('一级标题', 1)
doc.add_heading('二级标题', 2)
doc.add_heading('三级标题', 3)
doc.add_heading('四级标题', 4)

# 添加文本
para = doc.add_paragraph(text='这是文本内容！')
run = para.add_run(text='这是run中的内容')
run.font.size = Pt(32)
doc.save('test.docx')
