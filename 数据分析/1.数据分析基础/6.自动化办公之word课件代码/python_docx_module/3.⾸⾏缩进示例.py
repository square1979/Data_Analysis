from docx import Document
from docx.shared import Inches

doc = Document()
article = doc.add_paragraph()
a2 = article.add_run('因______，特向您请事假____天。请假时间⾃_____年___⽉___⽇⾄_____年___⽉___⽇。'
                     '这段时间内原计划安排的课程已做好处理，希望领导批准。')
# ⾸⾏缩进2个字符
article2_format = article.paragraph_format
article2_format.first_line_indent = Inches(0.3)
doc.save('test.docx')
