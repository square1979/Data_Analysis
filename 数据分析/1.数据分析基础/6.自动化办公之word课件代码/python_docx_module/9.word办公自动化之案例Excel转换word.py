import xlrd
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT


def read_excel():
    """读取Excel文件中的数据"""
    book = xlrd.open_workbook('test.xlsx')
    sheet = book.sheet_by_name(book.sheet_names()[0])
    nrows = sheet.nrows
    ncols = sheet.ncols
    datas = []
    keys = sheet.row_values(0)  # 读取标题行

    for row in range(1, nrows):
        data = {}
        first, second, third = '', '', ''
        for col in range(0, ncols):
            value = sheet.cell_value(row, col)
            value = value.replace('<', '').replace('>', '').replace('$', '')
            data[keys[col]] = value

            # 特殊处理第一列，按'/'分割
            if col == 0:
                arrs = value.lstrip('/').split('/')
                first, second, third = arrs[0], arrs[1] if len(arrs) > 1 else first, arrs[2] if len(
                    arrs) > 2 else second

        data['first'] = first
        data['second'] = second
        data['third'] = third
        datas.append(data)
    return datas


def write_word(datas):
    """生成Word文件"""
    if len(datas) < 1:
        print('Excel没有内容')
        return

        # 定义Word文档对象
    doc = Document()

    # 设置页面为横向
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    page_h, page_w = section.page_width, section.page_height
    section.page_width = page_w
    section.page_height = page_h

    # 设置字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 获取第三部分（部门）并去重
    data_third = []
    for data in datas:
        third = data['third']
        if third not in data_third:
            data_third.append(third)

            # 遍历部门并写入Word
    for third in data_third:
        # 写入部门，二级标题
        h2 = doc.add_heading(third, level=2)
        run = h2.runs[0]
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = '宋体'

        doc.add_paragraph(text='', style=None)  # 增加空行

        # 获取当前部门的模板
        data_template = []
        for data in datas:
            if data['third'] == third:
                template = {'first': data['first'], '模板名称': data['模板名称']}
                if template not in data_template:
                    data_template.append(template)

                    # 遍历模板
        for template in data_template:
            # 插入模板名称，三级标题
            h3 = doc.add_heading(template['模板名称'], level=3)
            run = h3.runs[0]
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = '宋体'

            doc.add_paragraph(text='', style=None)  # 换行

            # 筛选当前模板的数据
            data_table = [data for data in datas if
                          data['third'] == third and data['模板名称'] == template['模板名称'] and data['first'] ==
                          template['first']]

            # 新增表格
            doc_table = doc.add_table(rows=4, cols=5)
            doc_table.style = "Table Grid"
            doc_table.style.font.size = Pt(8)
            doc_table.style.font.name = '宋体'

            # 合并单元格并赋值
            doc_table.rows[0].cells[1].merge(doc_table.rows[0].cells[4])
            doc_table.rows[1].cells[1].merge(doc_table.rows[1].cells[4])
            doc_table.rows[2].cells[1].merge(doc_table.rows[2].cells[4])

            doc_table.rows[0].cells[0].text = '流程名称：'
            doc_table.rows[0].cells[1].text = data_table[0]['模板名称']
            doc_table.rows[1].cells[0].text = '使用人：'
            doc_table.rows[1].cells[1].text = data_table[0]['first']
            doc_table.rows[2].cells[0].text = '流程说明：'
            doc_table.rows[2].cells[1].text = data_table[0]['流程说明']

            # 设置表头
            head_cells = doc_table.rows[3].cells
            head_cells[0].text = '节点'
            head_cells[1].text = '节点名'
            head_cells[2].text = '处理⼈员'
            head_cells[3].text = '处理⽅式'
            head_cells[4].text = '跳转信息'
            # 设置列宽
            head_cells[0].width = Cm(1.9)
            head_cells[1].width = Cm(4.83)
            head_cells[2].width = Cm(8.25)
            head_cells[3].width = Cm(2.54)
            head_cells[4].width = Cm(5.64)
            # 第1 列⽔平居中，并设置⾏⾼，所有单元格内容垂直居中
            for i in range(0, 4):
                # ⽔平居中
                p = doc_table.rows[i].cells[0].paragraphs[0]
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                doc_table.rows[i].height = Cm(0.6)
                # ⾏⾼
                # 垂直居中
                for j in range(0, 5):
                    doc_table.rows[i].cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # ⽣成表格并填充内容
            row_num = 0
            for data in data_table:
                row = doc_table.add_row()
                row_cells = row.cells
                row_cells[0].text = str(row_num + 1)  # 序号，需要转换成字符串
                row_cells[1].text = data['节点名称']
                row_cells[2].text = data['审批⼈员']
                row_cells[3].text = data['审批⽅式']
                row_cells[4].text = ''
                # ⽔平居中
                p = row_cells[0].paragraphs[0]
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                row.height = Cm(0.6)  # ⾏⾼
                # 垂直居中
                for j in range(0, 5):
                    row_cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    row_num = row_num + 1
                    doc.add_paragraph(text='', style=None)  # 换⾏
                doc.save('test.docx')


print(write_word(read_excel()))
