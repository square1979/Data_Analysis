import docx.shared
from docx import Document
from docx.shared import Inches

doc = Document()
# 添加图片
doc.add_picture('logo.png', width=Inches(4), height=docx.shared.Cm(6))

# 增加分⻚
doc.add_page_break()
# 添加表格
table = doc.add_table(rows=3, cols=3, style='Table Grid')
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Name'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

# 再增加3⾏表格元素
for i in range(5):
    row_cells = table.add_row().cells
    row_cells[0].text = 'test' + str(i)
    row_cells[1].text = str(i)
    row_cells[2].text = 'desc' + str(i)

doc.save('test.docx')
