from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn

document = Document()
document.styles['Normal'].font.name = '宋体'
document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 标题设置
document.add_heading(u'MS WORD写⼊测试', 0)
document.add_heading('⼀级标题', 1)
document.add_heading('⼆级标题', 2)

# 写入段落
paragraph = document.add_paragraph('我们在做⽂本测试！')

# 设置样式
# 设置字号
run = paragraph.add_run('设置字号')
run.font.size = Pt(34)
# 设置字体
run = paragraph.add_run('设置中⽂字体')
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 设置字体倾斜
run = paragraph.add_run('字体倾斜')
run.italic = True
# 设置字体加粗
run = paragraph.add_run('字体加粗')
run.bold = True
# 设置列表元素顺序
document.add_paragraph('Intense quote', style='Intense Quote')
document.add_paragraph('⽆序列表元素1', style='List Bullet')
document.add_paragraph('⽆序列表元素3', style='List Bullet')
document.add_paragraph('有序列表元素1', style='List Number')
document.add_paragraph('有序列表元素3', style='List Number')
# 设置分页
document.add_page_break()
# 插入图片
document.add_picture('logo.png', width=Inches(1.35))
# 插入表格
table = document.add_table(rows=1, cols=3, style='Table Grid')
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Name'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
# 增加表格数据
for i in range(3):
    row_cells = table.add_row().cells
    row_cells[0].text = 'test' + str(i)
    row_cells[1].text = str(i)
    row_cells[2].text = 'desc' + str(i)

document.save('test.docx')
