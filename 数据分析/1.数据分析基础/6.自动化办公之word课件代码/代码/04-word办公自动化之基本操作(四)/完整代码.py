from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
document = Document()
document.add_heading(u'MS WORD写⼊测试',0)
document.add_heading('⼀级标题',1)
document.add_heading('⼆级标题',3)
paragraph = document.add_paragraph('我们在做⽂本测试！')
run = paragraph.add_run('设置字号、')
run.font.size = Pt(34)
run = paragraph.add_run('Set Font,')
run.font.name = 'Consolas'
run = paragraph.add_run('设置中⽂字体、')
run.font.name=u'宋体'
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run = paragraph.add_run('斜体、')
run.italic = True
paragraph.add_run('粗体').bold = True
document.add_paragraph('Intense quote',
style='Intense Quote')
document.add_paragraph(
 '⽆序列表元素1', style='List Bullet'
)
document.add_paragraph(
 '⽆序列表元素3', style='List Bullet'
)
document.add_paragraph(
 '有序列表元素1', style='List Number'
)
document.add_paragraph('⽣成效果有序列表元素3', style='List Number'
)
document.add_picture('dlrb.jpeg',
width=Inches(1.35))
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Name'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for i in range(3):
 row_cells = table.add_row().cells
 row_cells[0].text = 'test'+str(i)
 row_cells[1].text = str(i)
 row_cells[2].text = 'desc'+str(i)
document.add_page_break()
document.save('测试.docx')