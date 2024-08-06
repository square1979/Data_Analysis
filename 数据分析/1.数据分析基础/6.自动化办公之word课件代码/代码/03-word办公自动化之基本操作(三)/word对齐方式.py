from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
doc = Document()
title = doc.add_paragraph()
a2 = title.add_run('标题内容')
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # 居中
doc.save('test.docx')