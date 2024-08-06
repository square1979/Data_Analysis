from docx import Document
document = Document()
document.add_paragraph(
 '⽆序列表元素1', style='List Bullet'
 )
document.add_paragraph(
 '⽆序列表元素2', style='List Bullet'
 )
document.add_paragraph(
 '⽆序列表元素3', style='List Bullet'
 )

document.add_paragraph(
 '有序列表元素1', style='List Number'
 )
document.add_paragraph(
 '有序列表元素2', style='List Number'
 )
document.add_paragraph(
 '有序列表元素3', style='List Number'
 )


document.save('列表.docx')

