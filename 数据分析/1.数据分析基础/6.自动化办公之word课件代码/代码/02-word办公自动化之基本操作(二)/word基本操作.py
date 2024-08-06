from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

document = Document()

# 添加标题
document.add_heading('总标题', 0)
document.add_heading('一级标题', 1)
document.add_heading('二级标题', 2)

paragraph = document.add_paragraph('我是⽂本内容')
paragraph.add_run('粗体').bold = True

run = paragraph.add_run('设置字号、')
run.font.size = Pt(34)
run = paragraph.add_run('设置中⽂字体、')
run.font.name = '宋体'
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

run = paragraph.add_run('斜体、')
run.italic = True

document.save('test.docx')