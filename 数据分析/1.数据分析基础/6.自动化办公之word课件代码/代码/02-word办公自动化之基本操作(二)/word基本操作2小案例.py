from docx import Document
from docx.shared import Inches

doc = Document()
paragraph = doc.add_paragraph()
paragraph.add_run('因______，特向您请事假____天。请假时间自_____年___月___日至_____年___月___日。这段时间内原计划安排的课程已做好处理，希望领导批准。' )

article2_format = paragraph.paragraph_format
article2_format.first_line_indent = Inches(0.3)
doc.save('请假.docx')