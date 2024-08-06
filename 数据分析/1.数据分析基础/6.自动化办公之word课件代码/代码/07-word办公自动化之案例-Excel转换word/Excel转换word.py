import xlrd
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT


def read_excel():
    """读取Excel"""
    book = xlrd.open_workbook('test.xlsx')
    sheet = book.sheet_by_name(book.sheet_names()[0])
    nrows = sheet.nrows # 行数
    ncols = sheet.ncols # 列数
    datas = [] # 存放数据
    # 第一列 标题
    keys = sheet.row_values(0)
    for row in range(1, nrows):
      data = {} # 每一行数据
      for col in range(0, ncols):
        value = sheet.cell_value(row, col) # 取出每一个单元格的数据
        # 替换到特殊字符
        value = value.replace('<', '').replace('>', '').replace('$', '')
        data[keys[col]] = value
        # 截取第一列元素
        if col == 0:
          first = '' # 截取元素 第1
          second = '' # 截取元素 第2
          third = '' # 截取元素 第3
          arrs = value.lstrip('/').split('/') # 去掉第一个/ 然后再以/分组
          if len(arrs) > 0:
            if len(arrs) == 1:
              first = arrs[0]
              second = first
              third = second
            elif len(arrs) == 2:
              first = arrs[0]
              second = arrs[1]
              third = second
            elif len(arrs) == 3:
              first = arrs[0]
              second = arrs[1]
              third = arrs[2]
            else:
              first = arrs[0]
              second = arrs[1]
              third = arrs[2]
          else:
            first = value.ltrip('/')
            second = first
            third = second
          data['first'] = first
          data['second'] = second
          data['third'] = third
        # 截取第一列结束
      datas.append(data)
    return datas



def write_word(datas):
    """生成word文件"""
    if len(datas) < 1:
      print('Excel没有内容')
      return
    # 定义word文档对象
    doc = Document()
    # 添加横向
    section = doc.sections[0] # doc.add_section(start_type=WD_SECTION_START.CONTINUOUS) # 添加横向页的连续节
    section.orientation = WD_ORIENTATION.LANDSCAPE
    page_h, page_w = section.page_width, section.page_height
    section.page_width = page_w # 设置横向纸的宽度
    section.page_height = page_h # 设置横向纸的高度
    # 设置字体
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # 获取第3部分(部门) 并去重
    data_third = []
    for data in datas:
      third = data['third']
      if data_third.count(third) == 0:
        data_third.append(third)
    for third in data_third:
      h2 = doc.add_heading(third, level=2) # 写入部门,二级标题
      run = h2.runs[0] # 可以通过add_run来设置文字，也可以通过数组来获取
      run.font.color.rgb = RGBColor(0, 0, 0)
      run.font.name = u'宋体'
      doc.add_paragraph(text='', style=None) # 增加空白行 换行
      # 开始获取模板
      data_template = []
      for data in datas:
        if data['third'] == third:
          template = {'first': data['first'], '模板名称': data['模板名称']}
          if data_template.count(template) == 0:
            data_template.append(template)
      # 获取模板完成
      # 遍历模板
      for template in data_template:
        h3 = doc.add_heading(template['模板名称'], level=3) # 插入模板名称,三级标题
        run = h3.runs[0] # 可以通过add_run来设置文字，也可以通过数组来获取
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = u'宋体'
        doc.add_paragraph(text='', style=None) # 换行
        data_table = filter(
          lambda data: data['third'] == third and data['模板名称'] == template['模板名称'] and data['first'] ==
                 template['first'], datas)
        data_table = list(data_table)
        # 新增表格 4行5列
        doc_table = doc.add_table(rows=4, cols=5)
        doc_table.style = "Table Grid"
        doc_table.style.font.size = Pt(9)
        doc_table.style.font.name = '宋体'

        # 合并单元格 赋值
        doc_table.rows[0].cells[1].merge(doc_table.rows[0].cells[4])
        doc_table.rows[1].cells[1].merge(doc_table.rows[1].cells[4])
        doc_table.rows[2].cells[1].merge(doc_table.rows[2].cells[4])
        doc_table.rows[0].cells[0].text = '流程名称：'
        doc_table.rows[0].cells[1].text = data_table[0]['模板名称']
        doc_table.rows[1].cells[0].text = '使用人：'
        doc_table.rows[1].cells[1].text = data_table[0]['first']
        doc_table.rows[2].cells[0].text = '流程说明：'
        doc_table.rows[2].cells[1].text = data_table[0]['流程说明']

        # 设置标题
        head_cells = doc_table.rows[3].cells # 前面还有三行，特殊处理
        head_cells[0].text = '节点'
        head_cells[1].text = '节点名'
        head_cells[2].text = '处理人员'
        head_cells[3].text = '处理方式'
        head_cells[4].text = '跳转信息'
        # 设置列宽
        head_cells[0].width = Cm(1.9)
        head_cells[1].width = Cm(4.83)
        head_cells[2].width = Cm(8.25)
        head_cells[3].width = Cm(2.54)
        head_cells[4].width = Cm(5.64)
        # 第1 列水平居中，并设置行高，所有单元格内容垂直居中
        for i in range(0, 4):
          # 水平居中
          p = doc_table.rows[i].cells[0].paragraphs[0]
          p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
          doc_table.rows[i].height = Cm(0.6) # 行高
          # 垂直居中
          for j in range(0, 5):
            doc_table.rows[i].cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # 生成表格并填充内容
        row_num = 0
        for data in data_table:
          row = doc_table.add_row()
          row_cells = row.cells
          row_cells[0].text = str(row_num + 1) # 序号，需要转换成字符串
          row_cells[1].text = data['节点名称']
          row_cells[2].text = data['审批人员']
          row_cells[3].text = data['审批方式']
          row_cells[4].text = ''
          # 水平居中
          p = row_cells[0].paragraphs[0]
          p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
          row.height = Cm(0.6) # 行高
          # 垂直居中
          for j in range(0, 5):
            row_cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
          row_num = row_num + 1

        doc.add_paragraph(text='', style=None) # 换行
    doc.save('test.docx')


print(write_word(read_excel()))