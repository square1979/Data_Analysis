from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()
# 全局字体设置
doc.styles['Normal'].font.name = '宋体'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

# 创建标题
title = doc.add_paragraph()
title1 = title.add_run('公司员工请假条')
title1.font.size = Pt(20)
title1.blod = True
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 正文
article1 = doc.add_paragraph()
article2 = article1.add_run('___部：')

article3 = doc.add_paragraph()
article4 = article3.add_run('因____________ ，特向您请事假天。请假时间自_年月日 至_年月日。这段时间内原计 划安排的课程已做好处理，希望领导批准。')

# ⾸⾏缩进 负值表示悬挂缩进
article3_format = article3.paragraph_format
article3_format.first_line_indent = Inches(0.3)

article3 = doc.add_paragraph()
a3 = article3.add_run('请假人:') # 内容
article3.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
article3_format = article3.paragraph_format
article3_format.right_indent = Inches(0.9)

nowData = doc.add_paragraph()
n3 = nowData.add_run('年 月 日') # 内容
nowData.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
nowData_format = nowData.paragraph_format
nowData_format.right_indent = Inches(0.3)

doc.save('请假条.docx')