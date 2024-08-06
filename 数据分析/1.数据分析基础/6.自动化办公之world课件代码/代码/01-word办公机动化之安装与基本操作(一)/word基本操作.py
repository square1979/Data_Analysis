from docx import Document
from docx.shared import Pt

document = Document()

# 添加标题
document.add_heading('总标题', 0)
document.add_heading('一级标题', 1)
document.add_heading('二级标题', 2)

paragraph = document.add_paragraph('我是⽂本内容')

run = paragraph.add_run('设置字号、')
run.font.size = Pt(34)

document.save('test.docx')