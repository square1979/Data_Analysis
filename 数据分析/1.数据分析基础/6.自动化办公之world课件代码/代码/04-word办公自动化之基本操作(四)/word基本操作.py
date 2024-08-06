from docx import Document
from docx.shared import Inches

doc = Document()
# doc.add_picture('dlrb.jpeg', width=Inches(1.35))
doc.add_picture('dlrb.jpeg', width=Inches(1.35), height=Inches(1.35))

table = doc.add_table(rows=1, cols=3, style="Table Grid")
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Name'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for i in range(3):
    row_cells = table.add_row().cells
    row_cells[0].text = 'test'+str(i)
    row_cells[1].text = str(i)
    row_cells[2].text = 'desc'+str(i)
doc.save('test.docx')

