from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()
title = doc.add_paragraph()
a2 = title.add_run('标题内容')
# 左对⻬ LEFT;居中 CENTER;右对⻬ RIGHT;⽂本两端对⻬ JUSTIFY
title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# 增加列表
doc.add_paragraph('⽆序列表元素1', style='List Bullet')
doc.add_paragraph('⽆序列表元素2', style='List Bullet')
doc.add_paragraph('有序列表元素3', style='List Number')
doc.add_paragraph('有序列表元素4', style='List Number')

doc.save('test.docx')
