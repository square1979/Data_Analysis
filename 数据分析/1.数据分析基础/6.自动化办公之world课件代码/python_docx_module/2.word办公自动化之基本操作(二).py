from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches

# 新建对象
doc = Document()
# 写入段落
para = doc.add_paragraph()
# 写入run对象
run = para.add_run('设置中⽂字体')
# 设置字体样式
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(32)

# 设置字体样式——斜体
run1 = para.add_run('斜体')
run1.italic = True

# 设置字体样式——粗体
run2 = para.add_run('粗体')
run2.bold = True

# 首行缩进
article2_format = para.paragraph_format
article2_format.first_line_indent = Inches(2)


doc.save('test.docx')
